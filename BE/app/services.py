from docx import Document
import os
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

TEMPLATE_PATH = os.path.join(os.path.dirname(__file__), "doc_templates", "Contract-Template.docx")
EXPORT_FOLDER = os.path.join(os.path.dirname(__file__), "generated_docs")

def insert_paragraph_after(paragraph, text=None, style=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        new_para.add_run(text)
    if style is not None:
        new_para.style = style
    return new_para

def generate_contract(data: dict, export_filename: str):
    doc = Document(TEMPLATE_PATH)

    for i, paragraph in enumerate(doc.paragraphs):
        if "{{project_description}}" in paragraph.text:
            # Replace placeholder with heading text
            paragraph.text = "The Service Provider shall develop a working software with the following features:"
            insert_point = paragraph
            # Insert bullet content immediately after
            for section in data.get("project_description", []):
                insert_point = insert_paragraph_after(insert_point, f"• {section['title']}")
                for feature in section.get("features", []):
                    insert_point = insert_paragraph_after(insert_point, f"    - {feature}")
        else:
            for key, value in data.items():
                if key != "project_description":
                    placeholder = f"{{{{{key}}}}}"
                    if placeholder in paragraph.text:
                        paragraph.text = paragraph.text.replace(placeholder, str(value))

    os.makedirs(EXPORT_FOLDER, exist_ok=True)
    export_path = os.path.join(EXPORT_FOLDER, export_filename)
    doc.save(export_path)
    return export_path
